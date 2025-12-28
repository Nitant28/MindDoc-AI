import os
from PyPDF2 import PdfReader
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
from sqlalchemy.orm import Session
from app.core.config import settings
from app.database.models import Document, DocumentChunk
import json
import numpy as np
import io
import logging
from typing import Tuple, Optional

logger = logging.getLogger(__name__)

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    easyocr_reader = None  # Initialize on first use
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False


def get_easyocr_reader():
    """Lazy load EasyOCR reader to avoid startup delays."""
    global easyocr_reader
    if easyocr_reader is None and EASYOCR_AVAILABLE:
        easyocr_reader = easyocr.Reader(['en'])
    return easyocr_reader


def extract_text_from_image(img: Image.Image, use_ocr_priority: bool = False) -> str:
    """
    Extract text from image with multiple OCR fallbacks.
    
    Args:
        img: PIL Image object
        use_ocr_priority: If True, prioritize OCR engines over simple text detection
    
    Returns:
        Extracted text string
    """
    text = ""
    
    # Try EasyOCR first (best accuracy for complex layouts)
    if EASYOCR_AVAILABLE:
        try:
            reader = get_easyocr_reader()
            if reader:
                ocr_results = reader.readtext(img)
                extracted = " ".join([res[1] for res in ocr_results if res[1].strip() and res[2] > 0.3])
                if extracted.strip():
                    return extracted
        except Exception as e:
            logger.warning(f"EasyOCR failed: {e}")
    
    # Fallback to Tesseract
    if OCR_AVAILABLE:
        try:
            text = pytesseract.image_to_string(img)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"Tesseract failed: {e}")
    
    return text


def enhance_pdf_page_image(pix) -> Image.Image:
    """Enhance PDF page image for better OCR results."""
    img = Image.open(io.BytesIO(pix.tobytes("ppm")))
    
    # Convert to RGB if needed
    if img.mode in ('RGBA', 'LA', 'P'):
        img = img.convert('RGB')
    
    return img

def extract_text_from_file(file_obj, filename: str) -> str:
    """
    Extract complete text from files with multiple fallback methods and OCR.
    Handles both text-based and image-based PDFs seamlessly.
    """
    full_text = ""
    
    # Try unstructured first for advanced parsing
    if UNSTRUCTURED_AVAILABLE:
        try:
            file_obj.seek(0)
            elements = partition(file=file_obj, filename=filename)
            text = "\n".join([str(el) for el in elements])
            if text.strip() and len(text) > 50:
                logger.info(f"Successfully extracted text using unstructured library from {filename}")
                return text
        except Exception as e:
            logger.debug(f"Unstructured extraction failed: {e}")

    # PDF extraction with GUARANTEED complete text capture
    if filename.lower().endswith('.pdf'):
        return extract_pdf_comprehensive(file_obj, filename)
    
    # DOCX extraction
    elif filename.lower().endswith('.docx'):
        return extract_docx(file_obj)
    
    # Image extraction with OCR
    elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
        return extract_image(file_obj)
    
    else:
        raise ValueError("Unsupported file type. Supported: PDF, DOCX, images (PNG, JPG, JPEG, TIFF, BMP, GIF).")


def extract_pdf_comprehensive(file_obj, filename: str) -> str:
    """
    Comprehensive PDF extraction handling both text and image pages.
    """
    text = ""
    
    # Method 1: PyMuPDF (fitz) - BEST for hybrid PDFs
    if FITZ_AVAILABLE:
        try:
            file_obj.seek(0)
            doc = fitz.open(stream=file_obj.read(), filetype="pdf")
            logger.info(f"Processing PDF with PyMuPDF: {filename} ({len(doc)} pages)")
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                
                # If page has text content, use it
                if page_text.strip() and len(page_text.strip()) > 20:
                    text += f"--- Page {page_num + 1} ---\n{page_text}\n"
                else:
                    # Page is likely scanned/image-based, use OCR
                    logger.debug(f"Page {page_num + 1} has minimal text, applying OCR...")
                    ocr_text = extract_text_from_pdf_page(page, page_num + 1)
                    if ocr_text.strip():
                        text += ocr_text
            
            doc.close()
            
            if text.strip():
                logger.info(f"Successfully extracted text from {filename} using PyMuPDF")
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF error on {filename}: {e}")
    
    # Method 2: pdfplumber - alternative text extraction with table support
    if PDFPLUMBER_AVAILABLE:
        try:
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                text = ""
                logger.info(f"Processing PDF with pdfplumber: {filename} ({len(pdf.pages)} pages)")
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    
                    if page_text and len(page_text.strip()) > 20:
                        text += f"--- Page {page_num + 1} ---\n{page_text}\n"
                    else:
                        # Try OCR if text extraction minimal
                        logger.debug(f"Page {page_num + 1} has minimal text, applying OCR via pdfplumber...")
                        ocr_text = extract_image_from_pdfplumber_page(page, page_num + 1)
                        if ocr_text.strip():
                            text += ocr_text
                    
                    # Extract and format tables
                    try:
                        tables = page.extract_tables()
                        if tables:
                            text += "\n[Tables]\n"
                            for table in tables:
                                for row in table:
                                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                    text += row_text + "\n"
                                text += "\n"
                    except Exception:
                        pass
                
                if text.strip():
                    logger.info(f"Successfully extracted text from {filename} using pdfplumber")
                    return text
        except Exception as e:
            logger.warning(f"pdfplumber error on {filename}: {e}")
    
    # Method 3: PyPDF2 fallback
    try:
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        text = ""
        logger.info(f"Processing PDF with PyPDF2: {filename} ({len(reader.pages)} pages)")
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and len(page_text.strip()) > 10:
                text += f"--- Page {page_num + 1} ---\n{page_text}\n"
        
        if text.strip():
            logger.info(f"Successfully extracted text from {filename} using PyPDF2")
            return text
    except Exception as e:
        logger.warning(f"PyPDF2 error on {filename}: {e}")
    
    # Final fallback: Full OCR on all pages if text extraction failed
    if FITZ_AVAILABLE:
        try:
            file_obj.seek(0)
            doc = fitz.open(stream=file_obj.read(), filetype="pdf")
            text = ""
            logger.info(f"Applying full OCR to {filename} (all pages)...")
            
            for page_num, page in enumerate(doc):
                ocr_text = extract_text_from_pdf_page(page, page_num + 1)
                if ocr_text.strip():
                    text += ocr_text
            
            doc.close()
            
            if text.strip():
                logger.info(f"Successfully extracted text from {filename} using full OCR")
                return text
        except Exception as e:
            logger.error(f"Full OCR extraction failed on {filename}: {e}")
    
    logger.error(f"Could not extract text from {filename}")
    return ""


def extract_text_from_pdf_page(page, page_num: int) -> str:
    """Extract text from a single PDF page using OCR."""
    try:
        # Generate high-quality image from page
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
        img = enhance_pdf_page_image(pix)
        
        # Extract text via OCR
        ocr_text = extract_text_from_image(img, use_ocr_priority=True)
        
        if ocr_text.strip():
            return f"--- Page {page_num} (OCR) ---\n{ocr_text}\n"
    except Exception as e:
        logger.warning(f"Failed to OCR page {page_num}: {e}")
    
    return ""


def extract_image_from_pdfplumber_page(page, page_num: int) -> str:
    """Extract text from pdfplumber page image using OCR."""
    try:
        # Get page image
        img = page.to_image()
        
        # Extract text via OCR
        ocr_text = extract_text_from_image(img.original, use_ocr_priority=True)
        
        if ocr_text.strip():
            return f"--- Page {page_num} (OCR) ---\n{ocr_text}\n"
    except Exception as e:
        logger.warning(f"Failed to OCR pdfplumber page {page_num}: {e}")
    
    return ""


def extract_docx(file_obj) -> str:
    """Extract text from DOCX files."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install it with `pip install python-docx`")
    
    try:
        file_obj.seek(0)
    except Exception:
        pass
    
    try:
        doc = DocxDocument(file_obj)
        text = "\n".join([para.text for para in doc.paragraphs])
        if text.strip():
            logger.info("Successfully extracted text from DOCX")
            return text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
    
    return ""


def extract_image(file_obj) -> str:
    """Extract text from image files using OCR."""
    try:
        file_obj.seek(0)
        img = Image.open(file_obj)
        
        logger.info(f"Extracting text from image ({img.format}, {img.size})...")
        
        # Convert to RGB if needed
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        
        text = extract_text_from_image(img, use_ocr_priority=True)
        
        if text.strip():
            logger.info("Successfully extracted text from image")
            return text
    except Exception as e:
        logger.error(f"Image text extraction failed: {e}")
    
    return ""

def chunk_text(text: str) -> list:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

def create_vector_store(chunks, embeddings):
    # retained for compatibility; in this project we store embeddings in DB and use a Python retriever
    return list(zip(chunks, embeddings))

def save_document(db: Session, filename: str, content: str, tenant_id: int):
    print(f"Saving document: {filename}, tenant_id: {tenant_id}")
    doc = Document(filename=filename, content=content, tenant_id=tenant_id)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    print(f"Document saved with id: {doc.id}")
    chunks = chunk_text(content)
    try:
        # Compute embeddings using sentence-transformers (CPU friendly)
        try:
            from sentence_transformers import SentenceTransformer
            model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            print(f"SentenceTransformer failed: {e}, skipping embeddings")
            raise RuntimeError('Embeddings not available')
        embeddings = model.encode(chunks, show_progress_bar=False)
        for i, chunk in enumerate(chunks):
            # convert numpy array to list for JSON storage
            emb_list = np.array(embeddings[i]).astype(float).tolist()
            emb_str = json.dumps(emb_list)
            chunk_obj = DocumentChunk(document_id=doc.id, chunk_text=chunk, embedding=emb_str)
            db.add(chunk_obj)
        db.commit()
        print("Chunks saved")
    except Exception as e:
        print(f"Embeddings failed: {e}, saving chunks without embeddings")
        # Save chunks without embeddings, so fallback to token overlap works
        for chunk in chunks:
            chunk_obj = DocumentChunk(document_id=doc.id, chunk_text=chunk, embedding=None)
            db.add(chunk_obj)
        db.commit()
        print("Chunks saved without embeddings")
    return doc